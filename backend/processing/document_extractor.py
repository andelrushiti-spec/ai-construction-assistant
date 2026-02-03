"""
Universal document extractor for PDF, Word (DOCX/DOC), and Pages documents
"""
import logging
from pathlib import Path
from typing import List, Dict
import re
import zipfile
import tempfile
import shutil

# PDF support
import pdfplumber

# Word document support
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Word documents won't be supported.")

# Optional OCR support
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Only digital PDFs will be supported.")

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text from PDF, Word, and Pages documents"""

    def __init__(self, tesseract_cmd: str = None):
        """
        Initialize document extractor

        Args:
            tesseract_cmd: Path to tesseract executable (for OCR)
        """
        if tesseract_cmd and OCR_AVAILABLE:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    def extract_text(self, file_path: str, use_ocr: bool = False) -> List[Dict[str, any]]:
        """
        Extract text from document (auto-detects file type)

        Args:
            file_path: Path to document file
            use_ocr: Force OCR for PDFs

        Returns:
            List of dicts with page number and text: [{'page': 1, 'text': '...'}, ...]
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type by extension
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._extract_from_pdf(file_path, use_ocr)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_word(file_path)
        elif extension == '.pages':
            return self._extract_from_pages(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _extract_from_pdf(self, pdf_path: Path, use_ocr: bool = False) -> List[Dict[str, any]]:
        """Extract text from PDF using pdfplumber and optional OCR"""
        pages = []

        try:
            # First try pdfplumber for digital PDFs
            if not use_ocr:
                pages = self._extract_pdf_with_pdfplumber(pdf_path)

                # Check if extraction was successful
                if self._is_extraction_successful(pages):
                    logger.info(f"Successfully extracted text from {pdf_path.name} using pdfplumber")
                    return pages
                else:
                    logger.warning(f"pdfplumber extraction poor quality, falling back to OCR for {pdf_path.name}")

            # Fallback to OCR for scanned PDFs
            pages = self._extract_pdf_with_ocr(pdf_path)
            logger.info(f"Successfully extracted text from {pdf_path.name} using OCR")

        except Exception as e:
            logger.error(f"Error extracting text from {pdf_path.name}: {str(e)}")
            raise

        return pages

    def _extract_pdf_with_pdfplumber(self, pdf_path: Path) -> List[Dict[str, any]]:
        """Extract text using pdfplumber"""
        pages = []

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                pages.append({
                    'page': page_num,
                    'text': text.strip(),
                    'method': 'pdfplumber'
                })

        return pages

    def _extract_pdf_with_ocr(self, pdf_path: Path) -> List[Dict[str, any]]:
        """Extract text using Tesseract OCR"""
        if not OCR_AVAILABLE:
            logger.error("OCR not available. Please install tesseract-ocr and pytesseract.")
            raise RuntimeError("OCR libraries not installed. Please install tesseract-ocr, pytesseract, and pdf2image.")

        pages = []
        images = convert_from_path(str(pdf_path), dpi=300)

        for page_num, image in enumerate(images, start=1):
            text = pytesseract.image_to_string(image, lang='eng+sqi')
            pages.append({
                'page': page_num,
                'text': text.strip(),
                'method': 'ocr'
            })

        return pages

    def _extract_from_word(self, word_path: Path) -> List[Dict[str, any]]:
        """Extract text from Word document (.docx)"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed. Please install: pip install python-docx")

        logger.info(f"Extracting text from Word document: {word_path.name}")

        try:
            doc = DocxDocument(str(word_path))

            # Extract all paragraphs
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            all_text.append(cell.text.strip())

            # Combine all text
            full_text = '\n\n'.join(all_text)

            # Split into "pages" (chunks of ~500 words for consistency with PDF pages)
            pages = self._split_into_pages(full_text, words_per_page=500)

            logger.info(f"Extracted {len(pages)} pages from Word document")
            return pages

        except Exception as e:
            logger.error(f"Error extracting from Word document: {str(e)}")
            raise

    def _extract_from_pages(self, pages_path: Path) -> List[Dict[str, any]]:
        """Extract text from Apple Pages document (.pages)"""
        logger.info(f"Extracting text from Pages document: {pages_path.name}")

        try:
            # Pages files are actually zip archives containing a PDF
            # We can extract the PDF and process it
            temp_dir = tempfile.mkdtemp()

            try:
                with zipfile.ZipFile(str(pages_path), 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                # Look for preview.pdf or QuickLook/Preview.pdf
                pdf_paths = [
                    Path(temp_dir) / 'QuickLook' / 'Preview.pdf',
                    Path(temp_dir) / 'preview.pdf',
                ]

                for pdf_path in pdf_paths:
                    if pdf_path.exists():
                        logger.info(f"Found embedded PDF in Pages document: {pdf_path}")
                        pages = self._extract_from_pdf(pdf_path, use_ocr=False)

                        # Mark as Pages document
                        for page in pages:
                            page['method'] = 'pages'

                        return pages

                # If no PDF found, try extracting from index.xml
                index_xml = Path(temp_dir) / 'index.xml'
                if index_xml.exists():
                    logger.warning("No embedded PDF found, attempting XML extraction (limited)")
                    text = self._extract_from_pages_xml(index_xml)
                    return self._split_into_pages(text, words_per_page=500)

                raise ValueError("Could not find text content in Pages document")

            finally:
                # Clean up temp directory
                shutil.rmtree(temp_dir, ignore_errors=True)

        except zipfile.BadZipFile:
            # Older Pages files might not be zip archives
            logger.error("Pages document is not in a supported format (not a zip archive)")
            raise ValueError("Unsupported Pages document format. Please export as PDF or Word.")
        except Exception as e:
            logger.error(f"Error extracting from Pages document: {str(e)}")
            raise

    def _extract_from_pages_xml(self, xml_path: Path) -> str:
        """Extract text from Pages index.xml (basic extraction)"""
        try:
            with open(xml_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Very basic text extraction - remove XML tags
            # This is a fallback and may not capture all formatting
            text = re.sub(r'<[^>]+>', '', content)
            text = re.sub(r'\s+', ' ', text).strip()

            return text
        except Exception as e:
            logger.error(f"Error reading Pages XML: {str(e)}")
            return ""

    def _split_into_pages(self, text: str, words_per_page: int = 500) -> List[Dict[str, any]]:
        """Split continuous text into page-like chunks"""
        words = text.split()
        pages = []

        for i in range(0, len(words), words_per_page):
            page_words = words[i:i + words_per_page]
            page_text = ' '.join(page_words)

            pages.append({
                'page': len(pages) + 1,
                'text': page_text,
                'method': 'word_split'
            })

        return pages if pages else [{'page': 1, 'text': text, 'method': 'word_split'}]

    def _is_extraction_successful(self, pages: List[Dict[str, any]], min_chars_per_page: int = 50) -> bool:
        """Check if text extraction was successful"""
        if not pages:
            return False

        total_chars = sum(len(page['text']) for page in pages)
        avg_chars = total_chars / len(pages)

        return avg_chars >= min_chars_per_page

    def get_document_metadata(self, file_path: str) -> Dict[str, any]:
        """Get document metadata (page count, file size, type, etc.)"""
        file_path = Path(file_path)

        metadata = {
            'filename': file_path.name,
            'file_size': file_path.stat().st_size,
            'file_type': file_path.suffix.lower(),
            'pages': 0
        }

        try:
            extension = file_path.suffix.lower()

            if extension == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    metadata['pages'] = len(pdf.pages)
                    if pdf.metadata:
                        metadata['title'] = pdf.metadata.get('Title', '')
                        metadata['author'] = pdf.metadata.get('Author', '')

            elif extension == '.docx' and DOCX_AVAILABLE:
                doc = DocxDocument(str(file_path))
                # Estimate pages (roughly 500 words per page)
                total_words = sum(len(p.text.split()) for p in doc.paragraphs)
                metadata['pages'] = max(1, total_words // 500)

                # Try to get core properties
                try:
                    core_props = doc.core_properties
                    metadata['title'] = core_props.title or ''
                    metadata['author'] = core_props.author or ''
                except:
                    pass

            elif extension == '.pages':
                # Extract from Pages and count
                pages = self.extract_text(str(file_path))
                metadata['pages'] = len(pages)

        except Exception as e:
            logger.error(f"Error reading document metadata: {str(e)}")

        return metadata

    def extract_contract_metadata(self, pages: List[Dict[str, any]]) -> Dict[str, any]:
        """Extract contract-specific metadata from text"""
        combined_text = " ".join([p['text'] for p in pages[:3]])

        metadata = {
            'parties': [],
            'dates': [],
            'contract_number': None,
            'contract_type': None
        }

        # Extract contract number patterns
        contract_num_patterns = [
            r'Contract\s+(?:No\.?|Number)\s*[:.]?\s*([A-Z0-9/-]+)',
            r'Kontratë\s+(?:Nr\.?|Numër)\s*[:.]?\s*([A-Z0-9/-]+)',
        ]

        for pattern in contract_num_patterns:
            match = re.search(pattern, combined_text, re.IGNORECASE)
            if match:
                metadata['contract_number'] = match.group(1)
                break

        # Extract date patterns
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
            r'\d{1,2}\s+(?:Janar|Shkurt|Mars|Prill|Maj|Qershor|Korrik|Gusht|Shtator|Tetor|Nëntor|Dhjetor)\s+\d{4}',
        ]

        for pattern in date_patterns:
            dates = re.findall(pattern, combined_text, re.IGNORECASE)
            metadata['dates'].extend(dates)

        metadata['dates'] = list(set(metadata['dates']))[:5]

        return metadata

    def extract_law_metadata(self, pages: List[Dict[str, any]]) -> Dict[str, any]:
        """Extract law-specific metadata from text"""
        first_page_text = pages[0]['text'] if pages else ""

        metadata = {
            'law_number': None,
            'law_title': None,
            'year': None,
            'articles_count': 0
        }

        # Extract law number patterns
        law_num_patterns = [
            r'Law\s+(?:No\.?|Number)\s*[:.]?\s*([0-9/]+)',
            r'Ligj\s+(?:Nr\.?|Numër)\s*[:.]?\s*([0-9/]+)',
            r'VKM\s+(?:Nr\.?)\s*[:.]?\s*([0-9/]+)',
        ]

        for pattern in law_num_patterns:
            match = re.search(pattern, first_page_text, re.IGNORECASE)
            if match:
                metadata['law_number'] = match.group(1)
                break

        # Extract year
        year_match = re.search(r'\b(19|20)\d{2}\b', first_page_text)
        if year_match:
            metadata['year'] = int(year_match.group(0))

        # Count articles
        all_text = " ".join([p['text'] for p in pages])
        article_patterns = [
            r'Article\s+\d+',
            r'Neni\s+\d+',
        ]

        for pattern in article_patterns:
            articles = re.findall(pattern, all_text, re.IGNORECASE)
            metadata['articles_count'] = max(metadata['articles_count'], len(articles))

        return metadata
